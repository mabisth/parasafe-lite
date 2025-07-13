from fastapi import FastAPI, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, HttpUrl
import requests
import urllib3
from urllib.parse import urlparse, urljoin, urlunparse
import ssl
import socket
from datetime import datetime
import re
import json
from typing import List, Dict, Any, Optional
import os
from motor.motor_asyncio import AsyncIOMotorClient
import asyncio
import aiohttp
import time
from bs4 import BeautifulSoup
import warnings
import io
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# Disable SSL warnings for testing purposes
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
warnings.filterwarnings('ignore', message='Unverified HTTPS request')

app = FastAPI(title="ParaSafe-Lite API", description="Web Application Security Scanner")

# CORS configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# MongoDB setup
mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
client = AsyncIOMotorClient(mongo_url)
db = client.parasafe_lite

class ScanRequest(BaseModel):
    url: HttpUrl

class VulnerabilityReport(BaseModel):
    title: str
    description: str
    risk: str  # high, medium, low, info
    evidence: Optional[str] = None
    recommendation: str
    manual_verification: str
    owasp_category: Optional[str] = None
    cwe_id: Optional[str] = None

class ScanResult(BaseModel):
    scan_id: str
    target_url: str
    scan_time: str
    summary: Dict[str, int]
    vulnerabilities: List[VulnerabilityReport]
    scan_info: Dict[str, Any]

class ReportExporter:
    def __init__(self):
        pass

    def generate_pdf(self, scan_result: ScanResult) -> bytes:
        """Generate PDF report"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=72, bottomMargin=72)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        story = []
        
        # Title
        story.append(Paragraph("ParaSafe-Lite Security Report", title_style))
        story.append(Spacer(1, 20))
        
        # Scan Information
        story.append(Paragraph("Scan Information", styles['Heading2']))
        scan_info_data = [
            ['Target URL', scan_result.target_url],
            ['Scan Time', scan_result.scan_time],
            ['Scan ID', scan_result.scan_id],
            ['Server', scan_result.scan_info.get('server_info', 'Unknown')],
            ['Technologies', ', '.join(scan_result.scan_info.get('technologies', []))]
        ]
        
        scan_info_table = Table(scan_info_data, colWidths=[2*inch, 4*inch])
        scan_info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(scan_info_table)
        story.append(Spacer(1, 20))
        
        # Summary
        story.append(Paragraph("Vulnerability Summary", styles['Heading2']))
        summary_data = [
            ['Risk Level', 'Count'],
            ['High Risk', str(scan_result.summary.get('high', 0))],
            ['Medium Risk', str(scan_result.summary.get('medium', 0))],
            ['Low Risk', str(scan_result.summary.get('low', 0))],
            ['Informational', str(scan_result.summary.get('info', 0))]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(summary_table)
        story.append(PageBreak())
        
        # Vulnerabilities
        if scan_result.vulnerabilities:
            story.append(Paragraph("Vulnerability Details", styles['Heading2']))
            
            for i, vuln in enumerate(scan_result.vulnerabilities):
                story.append(Paragraph(f"{i+1}. {vuln.title}", styles['Heading3']))
                story.append(Paragraph(f"<b>Risk Level:</b> {vuln.risk.upper()}", styles['Normal']))
                story.append(Paragraph(f"<b>Description:</b> {vuln.description}", styles['Normal']))
                
                if vuln.evidence:
                    story.append(Paragraph(f"<b>Evidence:</b> {vuln.evidence}", styles['Normal']))
                
                story.append(Paragraph(f"<b>Recommendation:</b> {vuln.recommendation}", styles['Normal']))
                story.append(Paragraph(f"<b>Manual Verification:</b> {vuln.manual_verification}", styles['Normal']))
                
                if vuln.owasp_category:
                    story.append(Paragraph(f"<b>OWASP Category:</b> {vuln.owasp_category}", styles['Normal']))
                
                if vuln.cwe_id:
                    story.append(Paragraph(f"<b>CWE ID:</b> {vuln.cwe_id}", styles['Normal']))
                
                story.append(Spacer(1, 20))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_word(self, scan_result: ScanResult) -> bytes:
        """Generate Word document report"""
        doc = Document()
        
        # Title
        title = doc.add_heading('ParaSafe-Lite Security Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Scan Information
        doc.add_heading('Scan Information', level=1)
        scan_table = doc.add_table(rows=5, cols=2)
        scan_table.style = 'Table Grid'
        
        scan_data = [
            ('Target URL', scan_result.target_url),
            ('Scan Time', scan_result.scan_time),
            ('Scan ID', scan_result.scan_id),
            ('Server', scan_result.scan_info.get('server_info', 'Unknown')),
            ('Technologies', ', '.join(scan_result.scan_info.get('technologies', [])))
        ]
        
        for i, (key, value) in enumerate(scan_data):
            scan_table.cell(i, 0).text = key
            scan_table.cell(i, 1).text = str(value)
        
        # Summary
        doc.add_heading('Vulnerability Summary', level=1)
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('High Risk', str(scan_result.summary.get('high', 0))),
            ('Medium Risk', str(scan_result.summary.get('medium', 0))),
            ('Low Risk', str(scan_result.summary.get('low', 0))),
            ('Informational', str(scan_result.summary.get('info', 0))),
            ('Total', str(sum(scan_result.summary.values())))
        ]
        
        for i, (key, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = value
        
        # Vulnerabilities
        if scan_result.vulnerabilities:
            doc.add_page_break()
            doc.add_heading('Vulnerability Details', level=1)
            
            for i, vuln in enumerate(scan_result.vulnerabilities):
                doc.add_heading(f"{i+1}. {vuln.title}", level=2)
                
                p = doc.add_paragraph()
                p.add_run('Risk Level: ').bold = True
                p.add_run(vuln.risk.upper())
                
                p = doc.add_paragraph()
                p.add_run('Description: ').bold = True
                p.add_run(vuln.description)
                
                if vuln.evidence:
                    p = doc.add_paragraph()
                    p.add_run('Evidence: ').bold = True
                    p.add_run(vuln.evidence)
                
                p = doc.add_paragraph()
                p.add_run('Recommendation: ').bold = True
                p.add_run(vuln.recommendation)
                
                p = doc.add_paragraph()
                p.add_run('Manual Verification: ').bold = True
                p.add_run(vuln.manual_verification)
                
                if vuln.owasp_category:
                    p = doc.add_paragraph()
                    p.add_run('OWASP Category: ').bold = True
                    p.add_run(vuln.owasp_category)
                
                if vuln.cwe_id:
                    p = doc.add_paragraph()
                    p.add_run('CWE ID: ').bold = True
                    p.add_run(vuln.cwe_id)
                
                doc.add_paragraph()  # Add spacing
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_json(self, scan_result: ScanResult) -> dict:
        """Generate JSON report"""
        return {
            "report_type": "ParaSafe-Lite Security Report",
            "generated_at": datetime.now().isoformat(),
            "scan_data": scan_result.dict()
        }

class SecurityScanner:
    def __init__(self):
        self.session = requests.Session()
        self.session.verify = False  # For testing purposes
        self.session.timeout = 10
        self.session.headers.update({
            'User-Agent': 'ParaSafe-Lite Security Scanner 1.0'
        })

    def normalize_url(self, url: str) -> str:
        """Normalize URL for consistent processing"""
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        return url

    def get_base_info(self, url: str) -> Dict[str, Any]:
        """Get basic information about the target"""
        try:
            response = self.session.get(url, allow_redirects=True)
            
            # Parse HTML for technology detection
            soup = BeautifulSoup(response.text, 'html.parser')
            technologies = self.detect_technologies(response, soup)
            
            return {
                'status_code': response.status_code,
                'server': response.headers.get('Server', 'Unknown'),
                'final_url': response.url,
                'response_time': response.elapsed.total_seconds(),
                'technologies': technologies,
                'content_length': len(response.content),
                'headers': dict(response.headers)
            }
        except Exception as e:
            return {'error': str(e)}

    def detect_technologies(self, response, soup) -> List[str]:
        """Detect web technologies in use"""
        technologies = []
        
        # Check headers for technology indicators
        headers = response.headers
        
        if 'X-Powered-By' in headers:
            technologies.append(headers['X-Powered-By'])
        
        if 'Server' in headers:
            server = headers['Server'].lower()
            if 'apache' in server:
                technologies.append('Apache')
            elif 'nginx' in server:
                technologies.append('Nginx')
            elif 'iis' in server:
                technologies.append('IIS')
        
        # Check for frameworks in HTML
        html_content = str(soup).lower()
        
        # JavaScript frameworks
        if 'react' in html_content or 'reactjs' in html_content:
            technologies.append('React')
        if 'angular' in html_content:
            technologies.append('Angular')
        if 'vue' in html_content or 'vuejs' in html_content:
            technologies.append('Vue.js')
        if 'jquery' in html_content:
            technologies.append('jQuery')
        
        # CMS Detection
        if 'wp-content' in html_content or 'wordpress' in html_content:
            technologies.append('WordPress')
        if 'drupal' in html_content:
            technologies.append('Drupal')
        if 'joomla' in html_content:
            technologies.append('Joomla')
        
        return list(set(technologies))

    def check_https_security(self, url: str) -> List[VulnerabilityReport]:
        """Check HTTPS implementation and SSL/TLS security"""
        vulnerabilities = []
        parsed_url = urlparse(url)
        
        # Check if HTTPS is used
        if parsed_url.scheme == 'http':
            vulnerabilities.append(VulnerabilityReport(
                title="Insecure HTTP Protocol",
                description="The website is using HTTP instead of HTTPS, which means data is transmitted in plain text.",
                risk="high",
                evidence=f"URL scheme: {parsed_url.scheme}",
                recommendation="Implement HTTPS with a valid SSL/TLS certificate and redirect all HTTP traffic to HTTPS.",
                manual_verification="Try accessing the site with http:// and verify if it redirects to https://. Check if sensitive data (login forms, personal info) is transmitted over HTTP.",
                owasp_category="A02:2021 – Cryptographic Failures",
                cwe_id="CWE-319"
            ))
        else:
            # Check SSL/TLS configuration
            try:
                hostname = parsed_url.hostname
                context = ssl.create_default_context()
                
                with socket.create_connection((hostname, 443), timeout=10) as sock:
                    with context.wrap_socket(sock, server_hostname=hostname) as ssock:
                        cert = ssock.getpeercert()
                        ssl_version = ssock.version()
                        
                        # Check for weak SSL/TLS versions
                        if ssl_version in ['TLSv1', 'TLSv1.1', 'SSLv2', 'SSLv3']:
                            vulnerabilities.append(VulnerabilityReport(
                                title="Weak SSL/TLS Version",
                                description=f"The server supports weak SSL/TLS version: {ssl_version}",
                                risk="medium",
                                evidence=f"SSL/TLS version: {ssl_version}",
                                recommendation="Disable support for TLS 1.1 and below. Use TLS 1.2 or TLS 1.3.",
                                manual_verification="Use tools like SSL Labs SSL Test (ssllabs.com/ssltest/) to verify SSL/TLS configuration.",
                                owasp_category="A02:2021 – Cryptographic Failures",
                                cwe_id="CWE-326"
                            ))
                        
            except Exception:
                pass  # SSL check failed, but we don't want to flag this as a vulnerability without more info
        
        return vulnerabilities

    def check_security_headers(self, url: str) -> List[VulnerabilityReport]:
        """Check for missing security headers"""
        vulnerabilities = []
        
        try:
            response = self.session.get(url)
            headers = response.headers
            
            # Critical security headers to check
            security_headers = {
                'Strict-Transport-Security': {
                    'risk': 'medium',
                    'description': 'HTTP Strict Transport Security (HSTS) header is missing, which could allow protocol downgrade attacks.',
                    'recommendation': 'Add "Strict-Transport-Security: max-age=31536000; includeSubDomains" header.',
                    'owasp': 'A02:2021 – Cryptographic Failures',
                    'cwe': 'CWE-319'
                },
                'X-Content-Type-Options': {
                    'risk': 'low',
                    'description': 'X-Content-Type-Options header is missing, which could allow MIME type sniffing attacks.',
                    'recommendation': 'Add "X-Content-Type-Options: nosniff" header.',
                    'owasp': 'A05:2021 – Security Misconfiguration',
                    'cwe': 'CWE-79'
                },
                'X-Frame-Options': {
                    'risk': 'medium',
                    'description': 'X-Frame-Options header is missing, which could allow clickjacking attacks.',
                    'recommendation': 'Add "X-Frame-Options: DENY" or "X-Frame-Options: SAMEORIGIN" header.',
                    'owasp': 'A05:2021 – Security Misconfiguration',
                    'cwe': 'CWE-1021'
                },
                'Content-Security-Policy': {
                    'risk': 'medium',
                    'description': 'Content Security Policy (CSP) header is missing, which could allow XSS attacks.',
                    'recommendation': 'Implement a strict Content Security Policy header.',
                    'owasp': 'A03:2021 – Injection',
                    'cwe': 'CWE-79'
                },
                'X-XSS-Protection': {
                    'risk': 'low',
                    'description': 'X-XSS-Protection header is missing or disabled.',
                    'recommendation': 'Add "X-XSS-Protection: 1; mode=block" header.',
                    'owasp': 'A03:2021 – Injection',
                    'cwe': 'CWE-79'
                }
            }
            
            for header_name, header_info in security_headers.items():
                if header_name not in headers:
                    vulnerabilities.append(VulnerabilityReport(
                        title=f"Missing {header_name} Header",
                        description=header_info['description'],
                        risk=header_info['risk'],
                        evidence=f"Response headers: {', '.join(headers.keys())}",
                        recommendation=header_info['recommendation'],
                        manual_verification=f"Check response headers in browser dev tools or use curl -I {url} to verify the {header_name} header is present.",
                        owasp_category=header_info['owasp'],
                        cwe_id=header_info['cwe']
                    ))
            
        except Exception as e:
            pass  # Don't flag connection errors as vulnerabilities
        
        return vulnerabilities

    def check_information_disclosure(self, url: str) -> List[VulnerabilityReport]:
        """Check for information disclosure vulnerabilities"""
        vulnerabilities = []
        
        try:
            response = self.session.get(url)
            headers = response.headers
            content = response.text.lower()
            
            # Check for server information disclosure
            server_header = headers.get('Server', '')
            if server_header and any(version_indicator in server_header.lower() for version_indicator in ['/', 'apache', 'nginx', 'iis']):
                if re.search(r'\d+\.\d+', server_header):  # Has version number
                    vulnerabilities.append(VulnerabilityReport(
                        title="Server Version Information Disclosure",
                        description="The server header reveals version information that could help attackers identify vulnerabilities.",
                        risk="low",
                        evidence=f"Server header: {server_header}",
                        recommendation="Configure the web server to hide version information in response headers.",
                        manual_verification="Check the Server header in response using browser dev tools or curl -I command.",
                        owasp_category="A05:2021 – Security Misconfiguration",
                        cwe_id="CWE-200"
                    ))
            
            # Check for technology disclosure in HTML
            disclosure_patterns = [
                (r'generator.*wordpress \d+\.\d+', 'WordPress version disclosure'),
                (r'generator.*drupal \d+\.\d+', 'Drupal version disclosure'),
                (r'powered by.*php/\d+\.\d+', 'PHP version disclosure'),
            ]
            
            for pattern, title in disclosure_patterns:
                if re.search(pattern, content):
                    match = re.search(pattern, content)
                    vulnerabilities.append(VulnerabilityReport(
                        title=title,
                        description="Version information is disclosed in HTML content, which could help attackers identify vulnerabilities.",
                        risk="info",
                        evidence=f"Found in HTML: {match.group(0) if match else 'version pattern detected'}",
                        recommendation="Remove or hide version information from HTML output.",
                        manual_verification="View page source and search for version information in meta tags or comments.",
                        owasp_category="A05:2021 – Security Misconfiguration",
                        cwe_id="CWE-200"
                    ))
            
        except Exception:
            pass
        
        return vulnerabilities

    def check_common_vulnerabilities(self, url: str) -> List[VulnerabilityReport]:
        """Check for common web vulnerabilities"""
        vulnerabilities = []
        
        try:
            # Get the main page
            response = self.session.get(url)
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Check for forms (potential injection points)
            forms = soup.find_all('form')
            for form in forms:
                action = form.get('action', '')
                method = form.get('method', 'get').lower()
                
                # Check for forms without CSRF protection
                csrf_tokens = form.find_all(['input'], attrs={'name': re.compile(r'csrf|token|_token', re.I)})
                if not csrf_tokens and method == 'post':
                    vulnerabilities.append(VulnerabilityReport(
                        title="Missing CSRF Protection",
                        description="Form found without apparent CSRF token protection.",
                        risk="medium",
                        evidence=f"Form action: {action}, method: {method}",
                        recommendation="Implement CSRF tokens for all state-changing operations.",
                        manual_verification="Submit the form and check if it requires a CSRF token. Try submitting from a different origin.",
                        owasp_category="A01:2021 – Broken Access Control",
                        cwe_id="CWE-352"
                    ))
            
            # Check for admin/debug URLs
            admin_paths = ['/admin', '/administrator', '/admin.php', '/wp-admin', '/debug', '/test', '/.env']
            for path in admin_paths:
                try:
                    test_url = urljoin(url, path)
                    admin_response = self.session.get(test_url, timeout=5)
                    if admin_response.status_code == 200 and 'login' in admin_response.text.lower():
                        vulnerabilities.append(VulnerabilityReport(
                            title="Exposed Administrative Interface",
                            description=f"Administrative interface accessible at {path}",
                            risk="medium",
                            evidence=f"Status: {admin_response.status_code}, URL: {test_url}",
                            recommendation="Restrict access to administrative interfaces using IP whitelisting or VPN access.",
                            manual_verification=f"Navigate to {test_url} and verify if it's an admin panel accessible without proper authentication.",
                            owasp_category="A01:2021 – Broken Access Control",
                            cwe_id="CWE-287"
                        ))
                except:
                    continue
            
        except Exception:
            pass
        
        return vulnerabilities

    async def scan_website(self, url: str) -> ScanResult:
        """Perform comprehensive security scan"""
        scan_start = datetime.now()
        scan_id = f"scan_{int(time.time())}"
        
        # Normalize URL
        url = self.normalize_url(url)
        
        # Get basic information
        base_info = self.get_base_info(url)
        
        # Perform security checks
        vulnerabilities = []
        vulnerabilities.extend(self.check_https_security(url))
        vulnerabilities.extend(self.check_security_headers(url))
        vulnerabilities.extend(self.check_information_disclosure(url))
        vulnerabilities.extend(self.check_common_vulnerabilities(url))
        
        # Calculate summary
        summary = {
            'high': len([v for v in vulnerabilities if v.risk == 'high']),
            'medium': len([v for v in vulnerabilities if v.risk == 'medium']),
            'low': len([v for v in vulnerabilities if v.risk == 'low']),
            'info': len([v for v in vulnerabilities if v.risk == 'info'])
        }
        
        # Prepare scan info
        scan_info = {
            'target_url': url,
            'scan_time': scan_start.strftime('%Y-%m-%d %H:%M:%S UTC'),
            'scan_duration': (datetime.now() - scan_start).total_seconds(),
            'server_info': base_info.get('server', 'Unknown'),
            'technologies': base_info.get('technologies', []),
            'status_code': base_info.get('status_code'),
            'response_time': base_info.get('response_time')
        }
        
        # Create scan result
        result = ScanResult(
            scan_id=scan_id,
            target_url=url,
            scan_time=scan_start.strftime('%Y-%m-%d %H:%M:%S UTC'),
            summary=summary,
            vulnerabilities=vulnerabilities,
            scan_info=scan_info
        )
        
        # Store in database
        try:
            await db.scans.insert_one(result.dict())
        except Exception as e:
            print(f"Failed to store scan result: {e}")
        
        return result

# Initialize scanner and exporter
scanner = SecurityScanner()
exporter = ReportExporter()

@app.post("/api/scan", response_model=ScanResult)
async def scan_website(request: ScanRequest):
    """
    Scan a website for security vulnerabilities
    """
    try:
        url = str(request.url)
        result = await scanner.scan_website(url)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Scan failed: {str(e)}")

@app.post("/api/export/pdf")
async def export_pdf(scan_result: dict):
    """Export scan results as PDF"""
    try:
        result = ScanResult(**scan_result)
        pdf_bytes = exporter.generate_pdf(result)
        
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=parasafe-report-{result.scan_id}.pdf"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")

@app.post("/api/export/word")
async def export_word(scan_result: dict):
    """Export scan results as Word document"""
    try:
        result = ScanResult(**scan_result)
        word_bytes = exporter.generate_word(result)
        
        return StreamingResponse(
            io.BytesIO(word_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=parasafe-report-{result.scan_id}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word export failed: {str(e)}")

@app.post("/api/export/json")
async def export_json(scan_result: dict):
    """Export scan results as JSON"""
    try:
        result = ScanResult(**scan_result)
        json_data = exporter.generate_json(result)
        
        return JSONResponse(
            content=json_data,
            headers={"Content-Disposition": f"attachment; filename=parasafe-report-{result.scan_id}.json"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"JSON export failed: {str(e)}")

@app.get("/api/scans")
async def get_scan_history(limit: int = 10):
    """
    Get scan history
    """
    try:
        scans = await db.scans.find().sort("_id", -1).limit(limit).to_list(length=limit)
        return {"scans": scans}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to retrieve scans: {str(e)}")

@app.get("/api/scan/{scan_id}")
async def get_scan_result(scan_id: str):
    """
    Get specific scan result
    """
    try:
        scan = await db.scans.find_one({"scan_id": scan_id})
        if not scan:
            raise HTTPException(status_code=404, detail="Scan not found")
        return scan
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to retrieve scan: {str(e)}")

@app.get("/api/health")
async def health_check():
    """
    Health check endpoint
    """
    return {"status": "healthy", "service": "ParaSafe-Lite API", "version": "1.0.0"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)